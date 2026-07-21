"""Skenování a vyplňování šablon ve formátu .docx (Microsoft Word).

Používá python-docx. Placeholder může být ve Wordu rozdělený do více „runů“
(např. když do něj omylem zasáhne kontrola pravopisu) – engine to řeší tím,
že text placeholderu poskládá dohromady napříč runy a nahrazení vloží do
prvního z nich, takže se zachová okolní formátování odstavce.
"""
from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from .placeholder import PLACEHOLDER_RE, najdi_klice, normalizuj


def scan(cesta: str) -> list[str]:
    """Vrátí seznam klíčů placeholderů použitých v dokumentu."""
    doc = Document(cesta)
    text = "\n".join(_vsechen_text(doc))
    return najdi_klice(text)


def fill(cesta_sablony: str, hodnoty: dict[str, str], cesta_vystupu: str) -> str:
    """Vyplní šablonu hodnotami a uloží výsledek. Vrací cestu k výstupu."""
    doc = Document(cesta_sablony)
    for odstavec in _vsechny_odstavce(doc):
        _vypln_odstavec(odstavec, hodnoty)
    doc.save(cesta_vystupu)
    return cesta_vystupu


# ---- procházení dokumentu --------------------------------------------
def _vsechny_odstavce(doc: Document):
    """Všechny odstavce v těle, tabulkách, záhlaví i zápatí."""
    yield from _odstavce_v_kontejneru(doc)
    for sekce in doc.sections:
        yield from _odstavce_v_kontejneru(sekce.header)
        yield from _odstavce_v_kontejneru(sekce.footer)
        yield from _odstavce_v_kontejneru(sekce.first_page_header)
        yield from _odstavce_v_kontejneru(sekce.first_page_footer)


def _odstavce_v_kontejneru(kontejner):
    for odstavec in kontejner.paragraphs:
        yield odstavec
    for tabulka in kontejner.tables:
        for radek in tabulka.rows:
            for bunka in radek.cells:
                yield from _odstavce_v_kontejneru(bunka)


def _vsechen_text(doc: Document) -> list[str]:
    return [o.text for o in _vsechny_odstavce(doc)]


# ---- nahrazování v jednom odstavci -----------------------------------
def _vypln_odstavec(odstavec: Paragraph, hodnoty: dict[str, str]) -> None:
    runy = odstavec.runs
    if not runy:
        return
    # NFD → NFC po jednotlivých runech (diakritika v klíčích), indexy zůstanou sedět
    for r in runy:
        nfc = normalizuj(r.text)
        if nfc != r.text:
            r.text = nfc
    cely = "".join(r.text for r in runy)
    if "{{" not in cely:
        return

    # Nahrazuj od konce, ať se neposouvají indexy dřívějších výskytů.
    shody = list(PLACEHOLDER_RE.finditer(cely))
    if not shody:
        return
    for shoda in reversed(shody):
        klic = shoda.group(1)
        if klic not in hodnoty:
            continue
        _nahrad_rozsah(runy, shoda.start(), shoda.end(), hodnoty[klic])


def _nahrad_rozsah(runy, start: int, end: int, hodnota: str) -> None:
    """Nahradí znaky [start, end) v poskládaném textu runů hodnotou."""
    pozice = 0
    prvni_run = None
    offset_start = 0
    for run in runy:
        delka = len(run.text)
        konec = pozice + delka
        if prvni_run is None and start < konec:
            prvni_run = run
            offset_start = start - pozice
            zbytek_prvniho = run.text[:offset_start]
        if prvni_run is not None:
            if end <= konec:
                # poslední zasažený run
                zbytek_posledniho = run.text[end - pozice:]
                if run is prvni_run:
                    _nastav_text(run, zbytek_prvniho, hodnota, zbytek_posledniho)
                else:
                    _nastav_text(prvni_run, zbytek_prvniho, hodnota, "")
                    _nastav_text(run, "", "", zbytek_posledniho)
                break
            elif run is not prvni_run:
                run.text = ""  # celý run je uvnitř placeholderu
        pozice = konec


def _nastav_text(run, pred: str, hodnota: str, za: str) -> None:
    """Nastaví text runu jako ``pred + hodnota + za`` a víceřádkové hodnoty
    převede na skutečné zalomení řádku ve Wordu."""
    cely = pred + hodnota + za
    # vyčisti stávající <w:t>/<w:br> uzly
    for uzel in list(run._r):
        if uzel.tag in (qn("w:t"), qn("w:br")):
            run._r.remove(uzel)
    _pridej_text_s_zalomenim(run, cely)


def _pridej_text_s_zalomenim(run, text: str) -> None:
    casti = text.split("\n")
    for i, cast in enumerate(casti):
        if i > 0:
            run._r.append(run._r.makeelement(qn("w:br"), {}))
        t = run._r.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
        t.text = cast
        run._r.append(t)
