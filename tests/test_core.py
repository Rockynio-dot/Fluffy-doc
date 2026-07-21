"""Testy jádra: skenování placeholderů a vyplňování .docx i .odt."""
import os
import sys
import zipfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document

from protokoly.core import Generator, najdi_klice
from protokoly.core import docx_engine, odt_engine
from protokoly.models import Field, FieldType, Template


# ---- placeholder regex -----------------------------------------------
def test_najdi_klice_poradi_a_bez_duplicit():
    text = "SN: {{seriove_cislo}} model {{ model }} a znovu {{seriove_cislo}}"
    assert najdi_klice(text) == ["seriove_cislo", "model"]


def test_najdi_klice_s_diakritikou():
    text = "Značka: {{Značka_telefonu}}, jméno {{Jméno_Přebírajícího}}, {{Vybavení}}"
    assert najdi_klice(text) == ["Značka_telefonu", "Jméno_Přebírajícího", "Vybavení"]


def test_najdi_klice_nfd_rozlozena_diakritika():
    import unicodedata
    # jak to ukládá LibreOffice na některých systémech (rozložené znaky)
    text = unicodedata.normalize("NFD", "{{Značka_telefonu}} {{Vybavení}} {{Model}}")
    assert najdi_klice(text) == ["Značka_telefonu", "Vybavení", "Model"]


# ---- pomocníci --------------------------------------------------------
def _docx_se_sablonou(cesta):
    doc = Document()
    doc.add_paragraph("Sériové číslo: {{seriove_cislo}}")
    # placeholder rozdělený do více runů (simulace Wordu)
    p = doc.add_paragraph()
    p.add_run("Model: {{mo")
    p.add_run("del}} konec")
    doc.add_paragraph("Poznámka: {{poznamka}}")
    doc.save(cesta)


def _odt_se_sablonou(cesta):
    content = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"'
        ' xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" office:version="1.2">'
        "<office:body><office:text>"
        "<text:p>Sériové číslo: {{seriove_cislo}}</text:p>"
        "<text:p>Model: {{model}}</text:p>"
        "<text:p>Poznámka: {{poznamka}}</text:p>"
        "</office:text></office:body></office:document-content>"
    )
    with zipfile.ZipFile(cesta, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("mimetype", "application/vnd.oasis.opendocument.text", zipfile.ZIP_STORED)
        z.writestr("content.xml", content)


def _text_docx(cesta):
    return "\n".join(p.text for p in Document(cesta).paragraphs)


def _text_odt(cesta):
    with zipfile.ZipFile(cesta) as z:
        return z.read("content.xml").decode("utf-8")


# ---- docx -------------------------------------------------------------
def test_docx_scan_a_fill(tmp_path):
    sablona = tmp_path / "s.docx"
    _docx_se_sablonou(str(sablona))

    assert set(docx_engine.scan(str(sablona))) == {"seriove_cislo", "model", "poznamka"}

    vystup = tmp_path / "out.docx"
    docx_engine.fill(str(sablona), {
        "seriove_cislo": "SN-123",
        "model": "ThinkPad X1",     # rozdělený placeholder
        "poznamka": "řádek 1\nřádek 2",
    }, str(vystup))

    text = _text_docx(str(vystup))
    assert "SN-123" in text
    assert "ThinkPad X1 konec" in text     # okolní text zachován
    assert "{{" not in text
    assert "řádek 1" in text and "řádek 2" in text


# ---- odt --------------------------------------------------------------
def test_odt_scan_a_fill(tmp_path):
    sablona = tmp_path / "s.odt"
    _odt_se_sablonou(str(sablona))

    assert set(odt_engine.scan(str(sablona))) == {"seriove_cislo", "model", "poznamka"}

    vystup = tmp_path / "out.odt"
    odt_engine.fill(str(sablona), {
        "seriove_cislo": "SN-9 & spol.",   # XML-escaping
        "model": "Dell",
        "poznamka": "a\nb",
    }, str(vystup))

    xml = _text_odt(str(vystup))
    assert "SN-9 &amp; spol." in xml
    assert "<text:line-break/>" in xml
    assert "{{" not in xml
    # mimetype zůstal nekomprimovaný
    with zipfile.ZipFile(str(vystup)) as z:
        assert z.getinfo("mimetype").compress_type == zipfile.ZIP_STORED


# ---- generator + validace polí ---------------------------------------
def test_generator_validuje_povinne(tmp_path):
    sablona_doc = tmp_path / "s.docx"
    _docx_se_sablonou(str(sablona_doc))
    t = Template(nazev="Test", dokument="s.docx", slozka=str(tmp_path), pole=[
        Field("seriove_cislo", "S/N", FieldType.TEXT, povinne=True),
        Field("model", "Model", FieldType.TEXT),
        Field("poznamka", "Poznámka", FieldType.MULTILINE),
    ])

    import pytest
    with pytest.raises(Exception):
        Generator.generuj(t, {"model": "X"}, str(tmp_path / "o.docx"))

    out = Generator.generuj(t, {"seriove_cislo": "SN-1", "model": "X", "poznamka": ""},
                            str(tmp_path / "o.docx"))
    assert os.path.isfile(out)


def test_checkbox_formatovani():
    f = Field("souhlas", "Souhlas", FieldType.CHECKBOX)
    assert f.naformatuj("ano") == "☒ Ano"
    assert f.naformatuj("") == "☐ Ne"


# ---- ODT rozsekaný placeholder + diakritika --------------------------
def test_odt_rozsekany_placeholder_s_diakritikou(tmp_path):
    # LibreOffice rozdělil {{Značka_telefonu}} do dvou spanů
    content = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"'
        ' xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" office:version="1.2">'
        "<office:body><office:text>"
        "<text:p><text:span>Značka: {{Zna</text:span>"
        "<text:span>čka_telefonu}}</text:span></text:p>"
        "</office:text></office:body></office:document-content>"
    )
    sablona = tmp_path / "s.odt"
    with zipfile.ZipFile(sablona, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("mimetype", "application/vnd.oasis.opendocument.text", zipfile.ZIP_STORED)
        z.writestr("content.xml", content)

    assert odt_engine.scan(str(sablona)) == ["Značka_telefonu"]

    vystup = tmp_path / "out.odt"
    odt_engine.fill(str(sablona), {"Značka_telefonu": "Samsung"}, str(vystup))
    xml = _text_odt(str(vystup))
    assert "Samsung" in xml and "{{" not in xml


def test_odt_vnorene_spany(tmp_path):
    # přesně tak, jak placeholder rozseká LibreOffice (vnořené <text:span>)
    content = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"'
        ' xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" office:version="1.2">'
        "<office:body><office:text>"
        '<text:p>Značka: <text:span text:style-name="S">{{</text:span>'
        '<text:span text:style-name="S"><text:span text:style-name="T3">Značka_telefonu'
        '</text:span></text:span><text:span text:style-name="S">}}</text:span></text:p>'
        "</office:text></office:body></office:document-content>"
    )
    sablona = tmp_path / "s.odt"
    with zipfile.ZipFile(sablona, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("mimetype", "application/vnd.oasis.opendocument.text", zipfile.ZIP_STORED)
        z.writestr("content.xml", content)

    assert odt_engine.scan(str(sablona)) == ["Značka_telefonu"]
    vystup = tmp_path / "out.odt"
    odt_engine.fill(str(sablona), {"Značka_telefonu": "Samsung"}, str(vystup))
    xml = _text_odt(str(vystup))
    assert "Samsung" in xml and "{{" not in xml
    # výsledek musí zůstat validní XML (spany se nerozbily)
    import xml.dom.minidom as m
    m.parseString(xml)


def test_odt_nfd_scan_a_fill(tmp_path):
    import unicodedata
    content = unicodedata.normalize("NFD",
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"'
        ' xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" office:version="1.2">'
        "<office:body><office:text>"
        "<text:p>Značka: {{Značka_telefonu}}</text:p>"
        "<text:p>Vybavení: {{Vybavení}}</text:p>"
        "</office:text></office:body></office:document-content>"
    )
    sablona = tmp_path / "s.odt"
    with zipfile.ZipFile(sablona, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("mimetype", "application/vnd.oasis.opendocument.text", zipfile.ZIP_STORED)
        z.writestr("content.xml", content.encode("utf-8"))

    # klíče se najdou i z rozloženého (NFD) dokumentu, vrácené jsou v NFC
    assert odt_engine.scan(str(sablona)) == ["Značka_telefonu", "Vybavení"]

    vystup = tmp_path / "out.odt"
    odt_engine.fill(str(sablona), {"Značka_telefonu": "Samsung", "Vybavení": "Nabíječka"},
                    str(vystup))
    xml = _text_odt(str(vystup))
    assert "Samsung" in xml and "Nabíječka" in xml and "{{" not in xml


# ---- automatické číslo -----------------------------------------------
def test_auto_cislo_format_a_citac():
    from datetime import date
    f = Field("cislo_protokolu", "Číslo", FieldType.AUTO, vychozi="PST-{rok}-{poradi:04d}")
    assert f.format_auto(7, date(2026, 7, 21)) == "PST-2026-0007"
    assert f.validuj("") is None                 # auto se nevaliduje

    prazdny = Field("c", "C", FieldType.AUTO)     # bez vzoru -> 4místné číslo
    assert prazdny.format_auto(3) == "0003"


def test_template_pamatuje_citac_a_vzor(tmp_path):
    from protokoly.models.storage import Storage
    st = Storage(datovy_adresar=str(tmp_path))
    doc = tmp_path / "d.docx"
    Document().save(str(doc))
    t = Template(nazev="X", dokument="d.docx", citac=5, pole=[
        Field("cislo_protokolu", "Číslo", FieldType.AUTO, vychozi="A-{poradi:03d}"),
    ])
    st.uloz_sablonu(t, zdrojovy_dokument=str(doc))

    nactena = st.nacti_sablonu("x")
    assert nactena.citac == 5
    assert nactena.pole[0].vychozi == "A-{poradi:03d}"
    assert nactena.ma_auto_cislo()
