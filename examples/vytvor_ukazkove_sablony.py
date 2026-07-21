#!/usr/bin/env python3
"""Vytvoří ukázkové šablony předávacích protokolů (PC/NTB, telefon, SIM).

Spuštění::

    python examples/vytvor_ukazkove_sablony.py

Šablony se uloží do datové složky aplikace (``data/sablony``), takže je
appka rovnou nabídne. Slouží zároveň jako referenční příklad, jak vypadá
dokument se zástupnými poli ``{{klic}}``.
"""
from __future__ import annotations

import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from protokoly.models import Field, FieldType, Template
from protokoly.models.storage import Storage
from protokoly.core.placeholder import obal


# --------------------------------------------------------------------------
# definice ukázkových šablon: (název, popis, [Field], odt?)
# --------------------------------------------------------------------------
SPOLECNA_HLAVICKA = [
    Field("cislo_protokolu", "Číslo protokolu", FieldType.TEXT, povinne=True),
    Field("datum_predani", "Datum předání", FieldType.DATE, povinne=True),
    Field("predavajici", "Předávající", FieldType.TEXT, povinne=True),
    Field("prebirajici", "Přebírající", FieldType.TEXT, povinne=True),
    Field("stredisko", "Středisko / oddělení", FieldType.TEXT),
]

SABLONY = {
    "Předávací protokol – PC / Notebook": {
        "popis": "Předání počítače nebo notebooku zaměstnanci.",
        "pole": SPOLECNA_HLAVICKA + [
            Field("typ_zarizeni", "Typ zařízení", FieldType.CHOICE, moznosti=["PC", "Notebook"]),
            Field("vyrobce", "Výrobce", FieldType.TEXT),
            Field("model", "Model", FieldType.TEXT),
            Field("seriove_cislo", "Sériové číslo (S/N)", FieldType.TEXT, povinne=True),
            Field("inventarni_cislo", "Inventární číslo", FieldType.TEXT),
            Field("mac_adresa", "MAC adresa", FieldType.TEXT),
            Field("operacni_system", "Operační systém", FieldType.TEXT),
            Field("prislusenstvi", "Příslušenství", FieldType.MULTILINE,
                  vychozi="Napájecí adaptér, brašna"),
            Field("poznamka", "Poznámka", FieldType.MULTILINE),
        ],
    },
    "Předávací protokol – Mobilní telefon": {
        "popis": "Předání služebního mobilního telefonu.",
        "pole": SPOLECNA_HLAVICKA + [
            Field("vyrobce", "Výrobce", FieldType.TEXT),
            Field("model", "Model", FieldType.TEXT),
            Field("imei", "IMEI", FieldType.TEXT, povinne=True),
            Field("seriove_cislo", "Sériové číslo", FieldType.TEXT),
            Field("barva", "Barva", FieldType.TEXT),
            Field("prislusenstvi", "Příslušenství", FieldType.MULTILINE,
                  vychozi="Nabíječka, kabel USB-C"),
            Field("poznamka", "Poznámka", FieldType.MULTILINE),
        ],
    },
    "Předávací protokol – SIM karta": {
        "popis": "Předání služební SIM karty.",
        "pole": SPOLECNA_HLAVICKA + [
            Field("operator", "Operátor", FieldType.CHOICE,
                  moznosti=["T-Mobile", "O2", "Vodafone"]),
            Field("telefonni_cislo", "Telefonní číslo", FieldType.TEXT, povinne=True),
            Field("iccid", "ICCID (číslo SIM)", FieldType.TEXT, povinne=True),
            Field("puk", "PUK", FieldType.TEXT),
            Field("tarif", "Tarif", FieldType.TEXT),
            Field("poznamka", "Poznámka", FieldType.MULTILINE),
        ],
    },
}


def _postav_docx(nazev: str, popis: str, pole: list[Field], cesta: str) -> None:
    doc = Document()
    styl = doc.styles["Normal"]
    styl.font.name = "Calibri"
    styl.font.size = Pt(11)

    nadpis = doc.add_heading(nazev, level=0)
    nadpis.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if popis:
        p = doc.add_paragraph(popis)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True

    doc.add_paragraph()

    tabulka = doc.add_table(rows=0, cols=2)
    tabulka.style = "Table Grid"
    tabulka.columns[0].width = Pt(170)
    for f in pole:
        radek = tabulka.add_row().cells
        radek[0].paragraphs[0].add_run(f.popisek + (" *" if f.povinne else "")).bold = True
        radek[1].paragraphs[0].add_run(obal(f.klic))

    doc.add_paragraph()
    doc.add_paragraph(
        "Přebírající svým podpisem potvrzuje převzetí výše uvedeného majetku "
        "v pořádku a zavazuje se s ním zacházet dle interních směrnic."
    )
    doc.add_paragraph()

    podpisy = doc.add_table(rows=2, cols=2)
    podpisy.rows[0].cells[0].paragraphs[0].add_run("Předávající")
    podpisy.rows[0].cells[1].paragraphs[0].add_run("Přebírající")
    podpisy.rows[1].cells[0].paragraphs[0].add_run(obal("predavajici"))
    podpisy.rows[1].cells[1].paragraphs[0].add_run(obal("prebirajici"))

    doc.save(cesta)


def _postav_odt(nazev: str, popis: str, pole: list[Field], cesta: str) -> None:
    """Minimalní, ale validní ODT balíček (bez externí knihovny)."""
    import zipfile
    from xml.sax.saxutils import escape

    radky_xml = []
    for f in pole:
        radky_xml.append(
            "<table:table-row>"
            f"<table:table-cell table:style-name='cell'><text:p text:style-name='lbl'>"
            f"{escape(f.popisek)}{' *' if f.povinne else ''}</text:p></table:table-cell>"
            f"<table:table-cell table:style-name='cell'><text:p>{escape(obal(f.klic))}"
            "</text:p></table:table-cell>"
            "</table:table-row>"
        )
    tabulka = (
        "<table:table table:name='pole'>"
        "<table:table-column table:number-columns-repeated='2'/>"
        + "".join(radky_xml) +
        "</table:table>"
    )

    content = f"""<?xml version="1.0" encoding="UTF-8"?>
<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0" xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" office:version="1.2">
<office:automatic-styles>
<style:style style:name="lbl" style:family="paragraph"><style:text-properties fo:font-weight="bold"/></style:style>
<style:style style:name="cell" style:family="table-cell"><style:table-cell-properties fo:border="0.5pt solid #000000" fo:padding="0.1cm"/></style:style>
</office:automatic-styles>
<office:body><office:text>
<text:h text:outline-level="1">{escape(nazev)}</text:h>
<text:p>{escape(popis)}</text:p>
<text:p/>
{tabulka}
<text:p/>
<text:p>Přebírající svým podpisem potvrzuje převzetí výše uvedeného majetku v pořádku.</text:p>
<text:p/>
<text:p>Předávající: {escape(obal('predavajici'))}</text:p>
<text:p>Přebírající: {escape(obal('prebirajici'))}</text:p>
</office:text></office:body></office:document-content>"""

    styles = """<?xml version="1.0" encoding="UTF-8"?>
<office:document-styles xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" office:version="1.2"><office:styles/></office:document-styles>"""

    manifest = """<?xml version="1.0" encoding="UTF-8"?>
<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.2">
<manifest:file-entry manifest:full-path="/" manifest:media-type="application/vnd.oasis.opendocument.text"/>
<manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/>
<manifest:file-entry manifest:full-path="styles.xml" manifest:media-type="text/xml"/>
</manifest:manifest>"""

    with zipfile.ZipFile(cesta, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("mimetype", "application/vnd.oasis.opendocument.text", zipfile.ZIP_STORED)
        z.writestr("content.xml", content)
        z.writestr("styles.xml", styles)
        z.writestr("META-INF/manifest.xml", manifest)


def main() -> None:
    storage = Storage()
    with tempfile.TemporaryDirectory() as tmp:
        for i, (nazev, cfg) in enumerate(SABLONY.items()):
            pole = cfg["pole"]
            popis = cfg["popis"]
            # první šablonu vytvoříme i jako ODT, ať je vidět open formát
            pripona = ".odt" if i == 0 else ".docx"
            zdroj = os.path.join(tmp, f"sablona{i}{pripona}")
            if pripona == ".odt":
                _postav_odt(nazev, popis, pole, zdroj)
            else:
                _postav_docx(nazev, popis, pole, zdroj)

            sablona = Template(nazev=nazev, dokument="dokument" + pripona, popis=popis, pole=pole)
            storage.uloz_sablonu(sablona, zdrojovy_dokument=zdroj)
            print(f"  ✓ {nazev}  ({pripona})")

    print(f"\nHotovo. Šablony uloženy v: {storage.sablony_dir}")


if __name__ == "__main__":
    main()
